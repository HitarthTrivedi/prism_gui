#!/usr/bin/env python3
"""Turn the architecture Markdown into Word documents.

The Markdown under docs/architecture/ stays the source of truth — it diffs, it
reviews, and it lives next to the code it describes. This script produces the
Word copies that get mailed to people who will never open a repository.

    python docs/architecture/build_docx.py

Writes docs/architecture/word/:
    Prism-01-System-Overview.docx  …  Prism-08-State-and-Roadmap.docx
    Prism-Architecture-Complete.docx    (all nine, one file, one contents page)

What it handles: headings as REAL Word heading styles, so the navigation pane
and the contents field both work; pipe tables as real Word tables; fenced code
as shaded monospace; blockquotes as tinted callouts; and every ```mermaid block
rendered to PNG and embedded.

Diagrams need mermaid-cli, which is fetched on demand through npx:

    npx -y @mermaid-js/mermaid-cli@11

Without node the diagrams degrade to their source text in a captioned box
rather than failing the build — a document with a described diagram is worth
more than no document.

Fonts are deliberately Calibri / Consolas rather than the app's own Barlow:
Barlow is vendored into the build, not installed on a machine, and a .docx that
names a font the reader does not have silently falls back to something worse
than the safe choice would have been. The Industry accent (#5980a6) carries the
identity instead.
"""
from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
import tempfile

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(HERE, "word")
MERMAID_CFG = os.path.join(HERE, "mermaid-industry.json")

# ── the Industry palette, as the parts of it a Word file can carry ───────────
ACCENT = RGBColor(0x59, 0x80, 0xA6)
ACCENT_DEEP = RGBColor(0x3F, 0x62, 0x88)
INK = RGBColor(0x1C, 0x1F, 0x22)
INK_2 = RGBColor(0x4A, 0x50, 0x57)
INK_3 = RGBColor(0x76, 0x7C, 0x84)
SHADE_CODE = "F2F2F3"
SHADE_HEAD = "E4EBF2"
SHADE_QUOTE = "F4F6F9"
RULE = "C9CCD0"

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"

#: (filename, part number, title, subtitle). The part number drives the output
#: filename, the diagram captions and the cover, so it is stated once here
#: rather than parsed back out of the filename — "README.md".split("-")[0] is
#: "README.md", which is how you get Prism-README.md-Index.docx.
DOCS = [
    ("README.md", "00", "Index", "How to read this set"),
    ("01-system-overview.md", "01", "System Overview", "Layers, threading, startup, module map"),
    ("02-data-model.md", "02", "Data Model", "Every file, every field, the ER diagram"),
    ("03-data-flow.md", "03", "Data Flow", "Six pipelines, stage by stage"),
    ("04-api-reference.md", "04", "API Reference", "External services and the internal surface"),
    ("05-licensing.md", "05", "Licensing and Authorisation", "The gate on every launch"),
    ("06-addons.md", "06", "Add-on Subsystems", "Each add-on, and how well it is proven"),
    ("07-operations.md", "07", "Operations", "Build, CI, tests, failure catalogue"),
    ("08-state-and-roadmap.md", "08", "State and Roadmap", "Maturity board, triggers, debt"),
]


# ── low-level docx helpers ───────────────────────────────────────────────────

def shade(element, fill: str) -> None:
    """Paint a cell or paragraph background."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def para_borders(paragraph, *, left: str = "", bottom: str = "",
                 size: int = 6, space: int = 8) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for side, colour in (("left", left), ("bottom", bottom)):
        if not colour:
            continue
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), str(space))
        el.set(qn("w:color"), colour)
        borders.append(el)
    pPr.append(borders)


def keep_with_next(paragraph) -> None:
    """Stop a heading or caption being orphaned at the foot of a page."""
    pPr = paragraph._p.get_or_add_pPr()
    el = OxmlElement("w:keepNext")
    pPr.append(el)


def no_spacing(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def add_toc_field(doc) -> None:
    """A real contents field. Word offers to update it on open; the reader
    presses F9 or 'Update entire table' and gets live page numbers."""
    p = doc.add_paragraph()
    run = p.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and choose 'Update Field' to build the contents."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for el in (begin, instr, sep, placeholder, end):
        run._r.append(el)


def add_page_number_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.size = Pt(8.5)
    run.font.color.rgb = INK_3
    run.font.name = BODY_FONT

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, end):
        run._r.append(el)


# ── inline markdown ──────────────────────────────────────────────────────────

#: Order matters: code wins over everything (a backtick span is literal), then
#: bold, then links, then italics.
#:
#: The bold pattern is `\*\*.+?\*\*(?!\*)` rather than `\*\*[^*]+\*\*` because
#: these documents nest italics inside bold constantly — "**Both produced
#: *plausible* wrong numbers**" — and a class that excludes '*' simply fails to
#: match those, leaving the asterisks in the Word file as literal characters.
#: The trailing (?!\*) is what makes "*as well as***" close in the right place:
#: the first '**' of that '***' is the italic's closer, not the bold's.
_INLINE = re.compile(
    r"(?P<code>`[^`]+`)"
    r"|(?P<bold>\*\*.+?\*\*(?!\*))"
    r"|(?P<link>\[[^\]]+\]\([^)]+\))"
    r"|(?P<em>(?<![\*\w])\*[^*\n]+\*(?![\*\w]))"
)


def add_inline(paragraph, text: str, *, base_size: float = 10.5,
               colour: RGBColor | None = None, bold_all: bool = False) -> None:
    """Render one line of markdown into runs on `paragraph`."""
    text = text.replace("<br/>", "\n").replace("<br>", "\n")
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            _run(paragraph, text[pos:m.start()], base_size, colour, bold_all)
        kind = m.lastgroup
        raw = m.group()
        if kind == "code":
            r = _run(paragraph, raw[1:-1], base_size - 0.5, colour, bold_all)
            r.font.name = MONO_FONT
            r.font.color.rgb = ACCENT_DEEP
        elif kind == "bold":
            # Recurse: **`accounts_of(cfg)` is the one reader** is common in
            # these documents, and rendering the bold span as flat text left
            # the backticks sitting in the Word file as literal characters.
            # The bold pattern excludes '*', so this cannot recurse for ever.
            add_inline(paragraph, raw[2:-2], base_size=base_size,
                       colour=colour, bold_all=True)
        elif kind == "em":
            before = len(paragraph.runs)
            add_inline(paragraph, raw[1:-1], base_size=base_size,
                       colour=colour, bold_all=bold_all)
            for r in paragraph.runs[before:]:
                r.italic = True
        elif kind == "link":
            label = raw[1:raw.index("]")]
            r = _run(paragraph, label, base_size, colour, bold_all)
            r.font.color.rgb = ACCENT_DEEP
            r.underline = True
        pos = m.end()
    if pos < len(text):
        _run(paragraph, text[pos:], base_size, colour, bold_all)


def _run(paragraph, text: str, size: float, colour, bold: bool):
    # A newline inside a cell has to become a real line break, not a literal.
    parts = text.split("\n")
    run = None
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        run.font.size = Pt(size)
        run.font.name = BODY_FONT
        if colour is not None:
            run.font.color.rgb = colour
        if bold:
            run.bold = True
        if i < len(parts) - 1:
            run.add_break()
    return run


def strip_inline(text: str) -> str:
    """Markdown down to its words — for headings, where runs are styled once."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("`", "")
    return re.sub(r"(?<![\*\w])\*([^*\n]+)\*(?![\*\w])", r"\1", text).strip()


# ── mermaid ──────────────────────────────────────────────────────────────────

class Mermaid:
    """Renders mermaid blocks to PNG, once each, and remembers the failures."""

    def __init__(self, cache_dir: str):
        self.cache = cache_dir
        os.makedirs(cache_dir, exist_ok=True)
        self.available = shutil.which("npx") is not None
        self.rendered = 0
        self.failed = 0

    def png_for(self, source: str, key: str) -> str | None:
        if not self.available:
            return None
        out = os.path.join(self.cache, f"{key}.png")
        if os.path.exists(out):
            return out
        src = os.path.join(self.cache, f"{key}.mmd")
        with open(src, "w", encoding="utf-8") as f:
            f.write(source)
        cmd = ["npx", "-y", "@mermaid-js/mermaid-cli@11",
               "-i", src, "-o", out, "-b", "white", "-s", "2", "-w", "1500"]
        if os.path.exists(MERMAID_CFG):
            cmd += ["-c", MERMAID_CFG]
        try:
            proc = subprocess.run(cmd, capture_output=True, text=True,
                                  timeout=240, shell=(os.name == "nt"))
        except Exception as e:                                  # noqa: BLE001
            print(f"    ! {key}: {e}")
            self.failed += 1
            return None
        if proc.returncode != 0 or not os.path.exists(out):
            tail = (proc.stderr or proc.stdout or "").strip().splitlines()
            print(f"    ! {key}: {tail[-1] if tail else 'render failed'}")
            self.failed += 1
            return None
        self.rendered += 1
        return out


# ── the converter ────────────────────────────────────────────────────────────

NAV_LINE = re.compile(r"^\[←|^\[?←?\s*\[?Index\]|·\s*\[Next:|^\[← ")


def is_nav(line: str) -> bool:
    """The '← Index · Next →' lines are file navigation. Word has a contents
    page instead, and a dead relative link in a mailed document is noise."""
    s = line.strip()
    if not s.startswith("["):
        return False
    return ("Index](" in s or "Next:" in s or s.startswith("[←"))


#: A cell divider is an UNESCAPED pipe. Splitting on every '|' tore
#: "`FRESH` \| `GRACE` \| `STALE`" into three cells and shifted the whole row.
_CELL_SPLIT = re.compile(r"(?<!\\)\|")


def table_block(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Consume a pipe table starting at `start`. Returns (rows, next index)."""
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        raw = re.sub(r"^\|", "", raw)
        raw = re.sub(r"(?<!\\)\|$", "", raw)
        cells = [c.strip().replace(r"\|", "|") for c in _CELL_SPLIT.split(raw)]
        # the |---|:--:| separator carries alignment, not content
        if not all(re.fullmatch(r":?-{2,}:?", c or "-") for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for r, row in enumerate(rows):
        for c in range(width):
            cell = table.cell(r, c)
            cell.text = ""
            para = cell.paragraphs[0]
            no_spacing(para)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            text = row[c] if c < len(row) else ""
            if r == 0:
                shade(cell._tc.get_or_add_tcPr(), SHADE_HEAD)
                add_inline(para, text, base_size=9.5, colour=ACCENT_DEEP,
                           bold_all=True)
            else:
                add_inline(para, text, base_size=9.5)


def add_code(doc, code: str, *, caption: str = "") -> None:
    p = doc.add_paragraph()
    no_spacing(p)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.12)
    shade(p._p.get_or_add_pPr(), SHADE_CODE)
    para_borders(p, left=RULE, size=8, space=6)
    for i, line in enumerate(code.rstrip("\n").split("\n")):
        run = p.add_run(line)
        run.font.name = MONO_FONT
        run.font.size = Pt(8.5)
        run.font.color.rgb = INK
        if i < len(code.rstrip("\n").split("\n")) - 1:
            run.add_break()
    if caption:
        add_caption(doc, caption)


def add_caption(doc, text: str) -> None:
    p = doc.add_paragraph()
    no_spacing(p)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.name = BODY_FONT
    run.font.color.rgb = INK_3
    run.italic = True


def add_quote(doc, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    shade(p._p.get_or_add_pPr(), SHADE_QUOTE)
    para_borders(p, left="5980A6", size=18, space=10)
    add_inline(p, " ".join(lines), base_size=10.5, colour=INK_2)


#: Usable area inside the margins set in new_document(), in inches.
PORTRAIT_W, PORTRAIT_H = 6.3, 8.9
LANDSCAPE_W, LANDSCAPE_H = 9.2, 6.4
#: Past this the diagram is a horizontal strip and a portrait page wastes it.
WIDE_ASPECT = 1.6


def _orient(section, landscape: bool) -> None:
    """Set orientation AND swap the page dimensions — Word ignores the
    orientation flag on its own and keeps rendering the old page box."""
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    w, h = section.page_width, section.page_height
    if (landscape and w < h) or (not landscape and w > h):
        section.page_width, section.page_height = h, w


def image_size(png: str) -> tuple[float, float]:
    """(width, height) in inches for the page shape that suits this diagram."""
    try:
        from PIL import Image
        with Image.open(png) as im:
            px_w, px_h = im.size
    except Exception:                                           # noqa: BLE001
        return PORTRAIT_W, 0.0
    aspect = px_w / px_h if px_h else 1.0
    box_w, box_h = (LANDSCAPE_W, LANDSCAPE_H) if aspect > WIDE_ASPECT \
        else (PORTRAIT_W, PORTRAIT_H)
    width = box_w
    height = width / aspect
    if height > box_h:                       # tall: give it the full page height
        height = box_h
        width = height * aspect
    return width, height


def add_diagram(doc, png: str | None, source: str, caption: str) -> None:
    """Place one diagram, on the page shape it actually needs.

    A 2.3:1 entity-relationship diagram squeezed into a 6.3in portrait column
    is 2.7in tall and its labels are unreadable on paper. Wide diagrams get a
    landscape page of their own; tall ones get the full portrait height.
    """
    if not (png and os.path.exists(png)):
        add_code(doc, source, caption=f"{caption} — diagram source "
                                      "(mermaid renderer unavailable)")
        return

    width_in, height_in = image_size(png)
    landscape = width_in > PORTRAIT_W + 0.01

    if landscape:
        sect = doc.add_section(WD_SECTION.NEW_PAGE)
        _orient(sect, True)

    p = doc.add_paragraph()
    no_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.add_run().add_picture(png, width=Inches(width_in))
    add_caption(doc, caption)

    if landscape:
        back = doc.add_section(WD_SECTION.NEW_PAGE)
        _orient(back, False)


def convert(md_path: str, doc, mm: Mermaid, part: str,
            *, heading_offset: int = 0, title: str = "") -> None:
    """Append one markdown file to `doc`."""
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    diagram_n = 0
    i = 0
    pending_list_num = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── skip file-navigation lines and horizontal rules ──
        if not stripped or stripped in ("---", "***", "___"):
            i += 1
            continue
        if is_nav(stripped):
            i += 1
            continue

        # ── fenced blocks ──
        if stripped.startswith("```"):
            lang = stripped[3:].strip().lower()
            i += 1
            buf = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            body = "\n".join(buf)
            if lang == "mermaid":
                diagram_n += 1
                key = f"{part}-{diagram_n:02d}"
                caption = f"Diagram {int(part)}.{diagram_n}"
                add_diagram(doc, mm.png_for(body, key), body, caption)
            else:
                add_code(doc, body)
            continue

        # ── headings ──
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = strip_inline(stripped[level:])
            # the file's own H1 is replaced by the section title we were given
            if level == 1 and title:
                text = title
            level = min(level + heading_offset, 4)
            h = doc.add_heading(text, level=max(1, level))
            keep_with_next(h)
            for run in h.runs:
                run.font.name = BODY_FONT
                run.font.color.rgb = ACCENT_DEEP if level <= 2 else ACCENT
            i += 1
            continue

        # ── tables ──
        if stripped.startswith("|"):
            rows, i = table_block(lines, i)
            add_table(doc, rows)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue

        # ── blockquotes ──
        if stripped.startswith(">"):
            buf = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            add_quote(doc, [b for b in buf if b])
            continue

        # ── lists ──
        m_ul = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        m_ol = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m_ul or m_ol:
            indent = len((m_ul or m_ol).group(1))
            body = [m_ul.group(2) if m_ul else m_ol.group(3)]
            i += 1
            # A wrapped list item is one item, not two paragraphs. Continuation
            # lines are indented and start no new block.
            while i < len(lines):
                nxt = lines[i]
                s = nxt.strip()
                if (not s or s.startswith(("#", "|", ">", "```", "---"))
                        or re.match(r"^\s*([-*+]|\d+\.)\s+", nxt)
                        or is_nav(s) or not nxt.startswith((" ", "\t"))):
                    break
                body.append(s)
                i += 1
            style = "List Bullet" if m_ul else "List Number"
            if indent >= 2:
                style += " 2"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph()
            no_spacing(p)
            p.paragraph_format.space_after = Pt(3)
            add_inline(p, " ".join(body))
            continue

        # ── ordinary paragraph: gather until a blank or a block starts ──
        buf = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith(("#", "|", ">", "```", "---"))
                    or re.match(r"^\s*([-*+]|\d+\.)\s+", lines[i])
                    or is_nav(nxt)):
                break
            buf.append(nxt)
            i += 1
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        add_inline(p, " ".join(buf))


# ── document scaffolding ─────────────────────────────────────────────────────

def new_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)
    style.font.color.rgb = INK
    style.paragraph_format.space_after = Pt(7)
    style.paragraph_format.line_spacing = 1.14

    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.95)
        section.right_margin = Inches(0.95)
        add_page_number_footer(section)
    return doc


def title_page(doc, title: str, subtitle: str, *, part: str = "") -> None:
    for _ in range(3):
        doc.add_paragraph()

    if part:
        p = doc.add_paragraph()
        no_spacing(p)
        r = p.add_run(part.upper())
        r.font.size = Pt(10)
        r.font.name = BODY_FONT
        r.font.color.rgb = ACCENT
        r.bold = True

    p = doc.add_paragraph()
    no_spacing(p)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.font.size = Pt(34)
    r.font.name = BODY_FONT
    r.font.color.rgb = INK
    r.bold = True

    p = doc.add_paragraph()
    para_borders(p, bottom=RULE, size=8, space=10)
    r = p.add_run(subtitle)
    r.font.size = Pt(13)
    r.font.name = BODY_FONT
    r.font.color.rgb = INK_2

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(16)
    for label, value in (("Product", "Prism  ·  version 1.3.0"),
                         ("Commit", "20292f4"),
                         ("Drawn", "26 August 2026"),
                         ("Source", "docs/architecture/ — Markdown is the master")):
        r = meta.add_run(f"{label}\t")
        r.font.size = Pt(9)
        r.font.name = BODY_FONT
        r.font.color.rgb = INK_3
        r.bold = True
        r = meta.add_run(f"{value}\n")
        r.font.size = Pt(9)
        r.font.name = BODY_FONT
        r.font.color.rgb = INK_2


def main() -> int:
    os.makedirs(OUT_DIR, exist_ok=True)
    cache = os.path.join(tempfile.gettempdir(), "prism-mermaid-cache")
    mm = Mermaid(cache)
    if not mm.available:
        print("! npx not found — diagrams will appear as source text.")

    made = []

    # ── one file per document ──
    for filename, part, title, subtitle in DOCS:
        path = os.path.join(HERE, filename)
        if not os.path.exists(path):
            print(f"! missing {filename}")
            continue
        print(f"  {filename} -> {title}")
        doc = new_document()
        title_page(doc, title, subtitle,
                   part="Prism architecture" if part == "00" else
                        f"Prism architecture · part {part}")
        doc.add_page_break()
        convert(path, doc, mm, part, title=title)
        safe = re.sub(r"[^A-Za-z0-9]+", "-", title).strip("-")
        out = os.path.join(OUT_DIR, f"Prism-{part}-{safe}.docx")
        doc.save(out)
        made.append(out)

    # ── and one file with all of them ──
    print("  building the combined document")
    book = new_document()
    title_page(book, "Prism Architecture",
               "The complete reference — system, data, flow, operations",
               part="Alphakore · internal")
    book.add_page_break()

    h = book.add_heading("Contents", level=1)
    for run in h.runs:
        run.font.name = BODY_FONT
        run.font.color.rgb = ACCENT_DEEP
    add_toc_field(book)
    book.add_page_break()

    for filename, part, title, subtitle in DOCS:
        path = os.path.join(HERE, filename)
        if not os.path.exists(path):
            continue
        convert(path, book, mm, part, title=title)
        book.add_page_break()

    out = os.path.join(OUT_DIR, "Prism-Architecture-Complete.docx")
    book.save(out)
    made.append(out)

    print(f"\n{len(made)} documents in {OUT_DIR}")
    print(f"diagrams: {mm.rendered} rendered, {mm.failed} fell back to source")
    return 0


if __name__ == "__main__":
    sys.exit(main())
